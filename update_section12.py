"""
update_section12.py
===================
Replaces Section 12 in DXY_Pair_Strategy_Manual.docx with a comprehensive
updated version covering:
  - Two-indicator webhook architecture (DXYTradeAlert + DXYPairLevels)
  - DXY_ENTRY and DXY_EXIT JSON schemas (new in DXYTradeAlert.pine)
  - ENTRY, EXIT, SL_MOVE JSON schemas (DXYPairLevels.pine, SL_MOVE is new)
  - TradingView alert setup for both indicators
  - Automation routing logic
  - Platform guidance table
  - Google Sheets via Make example
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE    = Path(r"C:\Users\justi\OneDrive\Documents\Claude\Projects\DXY Backtesting")
DOCPATH = BASE / "DXY_Pair_Strategy_Manual.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
# Stored as (r, g, b) tuples for XML hex operations; converted to RGBColor for font use.
NAVY   = (0x0D, 0x2B, 0x5A)
GOLD   = (0xC9, 0xA0, 0x2C)
WHITE  = (0xFF, 0xFF, 0xFF)
BLACK  = (0x00, 0x00, 0x00)
LGREY  = (0xF2, 0xF2, 0xF2)
DKGREY = (0x60, 0x60, 0x60)
GREEN  = (0x1E, 0x6B, 0x1E)
LGREEN = (0xE2, 0xF0, 0xE2)
AMBER  = (0xB8, 0x5C, 0x00)
LAMBER = (0xFF, 0xF0, 0xD4)
LBLUE  = (0xD0, 0xE4, 0xF7)
DKBLUE = (0x1A, 0x3A, 0x6B)
DKRED  = (0xCC, 0x33, 0x00)

def rgb(t):
    """Convert (r,g,b) tuple to RGBColor for font.color.rgb assignments."""
    return RGBColor(*t)

def hex6(t) -> str:
    """Convert (r,g,b) tuple to 6-char uppercase hex string for XML."""
    return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, colour_tuple):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6(colour_tuple))
    tcPr.append(shd)


def set_cell_border(cell, sides=("top","bottom","left","right"), color="C9A02C", sz=8):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def hdr_para(doc, text, level=1):
    """Add a numbered section heading."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"] if level == 1 else doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = rgb(NAVY)
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    return p


def sub_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = rgb(NAVY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.bold   = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = rgb(colour)
    p.paragraph_format.space_after = Pt(4)
    return p


def code_block(doc, lines):
    """Monospaced JSON / code example block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(DKGREY)
    # light grey background via paragraph shading
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F4F4F4")
    pPr.append(shd)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Build a styled table with navy header and alternating rows."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, NAVY)
        set_cell_border(cell, color="C9A02C")
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9)
        run.font.color.rgb = rgb(WHITE)
        p.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        bg = LGREY if ri % 2 == 0 else WHITE
        tr = tbl.rows[ri + 1]
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="CCCCCC", sz=4)
            p   = cell.paragraphs[0]
            if isinstance(cell_text, tuple):
                # (text, bold, colour)
                run = p.add_run(cell_text[0])
                run.font.bold  = cell_text[1] if len(cell_text) > 1 else False
                run.font.color.rgb = rgb(cell_text[2]) if len(cell_text) > 2 else rgb(BLACK)
            else:
                run = p.add_run(str(cell_text))
            run.font.size  = Pt(9)

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()   # spacer
    return tbl


# ─────────────────────────────────────────────────────────────────────────────
#  Find and delete existing Section 12 content
# ─────────────────────────────────────────────────────────────────────────────

doc = Document(str(DOCPATH))

# Locate paragraphs that belong to Section 12 (between heading "12." and "13.")
body_el = doc.element.body
all_paras = doc.paragraphs   # includes tables? No — paragraphs only
all_elements = list(body_el)  # includes both <w:p> and <w:tbl>

# Find the index of the first element that starts section 12 heading
sec12_start = None
sec13_start = None

for i, elem in enumerate(all_elements):
    if elem.tag.endswith("}p"):
        # Check paragraph text
        text = "".join(r.text for r in elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
        stripped = text.strip()
        if stripped.startswith("12.") or stripped.startswith("12 "):
            if sec12_start is None:
                sec12_start = i
        if (stripped.startswith("13.") or stripped.startswith("13 ")) and sec12_start is not None:
            if sec13_start is None:
                sec13_start = i
                break

if sec12_start is None:
    raise RuntimeError("Could not find Section 12 heading in the document")

print(f"Section 12 starts at element index {sec12_start}")
print(f"Section 13 starts at element index {sec13_start}")

# Remove all elements from sec12_start up to (but not including) sec13_start
elements_to_remove = all_elements[sec12_start:sec13_start]
for elem in elements_to_remove:
    body_el.remove(elem)
print(f"Removed {len(elements_to_remove)} elements from Section 12")

# Find the new position of Section 13 (now shifted)
all_elements_new = list(body_el)
sec13_new_idx = None
for i, elem in enumerate(all_elements_new):
    if elem.tag.endswith("}p"):
        text = "".join(r.text for r in elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
        stripped = text.strip()
        if stripped.startswith("13.") or stripped.startswith("13 "):
            sec13_new_idx = i
            break

print(f"Section 13 is now at element index {sec13_new_idx}")

# ─────────────────────────────────────────────────────────────────────────────
#  Build a temporary document with the new Section 12 content
# ─────────────────────────────────────────────────────────────────────────────

tmp = Document()

# ── Section 12 heading ────────────────────────────────────────────────────────
p = tmp.add_paragraph()
run = p.add_run("12.  Alert, Webhook & Trade Journal Integration")
run.font.bold  = True
run.font.size  = Pt(14)
run.font.color.rgb = rgb(NAVY)

# ── 12.1 Architecture overview ────────────────────────────────────────────────
body(tmp, "The DXY strategy uses two TradingView indicators that each fire webhooks with "
     "distinct but complementary roles. Both must be enabled to get full automation coverage.")

make_table(tmp,
    headers=["Indicator", "Chart", "Webhook Role", "Events Fired"],
    rows=[
        ("DXYTradeAlert.pine", "DXY (index) chart", "Control channel — DXY-level signal and exit events", "DXY_ENTRY, DXY_EXIT"),
        ("DXYPairLevels.pine", "Each pair chart ×8", "Trade channel — per-pair entry, exit, and management", "ENTRY, EXIT, SL_MOVE"),
    ],
    col_widths=[4.5, 3.5, 6.0, 4.0]
)

body(tmp,
     "Typical live setup: 9 TradingView chart tabs open simultaneously "
     "(1 DXY + 8 pairs). Each has one 'alert() function calls' alert pointing to the same "
     "webhook endpoint. Your automation receives all events and routes by the 'event' field.")

# ── 12.2 DXYTradeAlert.pine webhooks (NEW) ───────────────────────────────────
sub_heading(tmp, "12.2  DXYTradeAlert.pine — DXY_ENTRY and DXY_EXIT events")

body(tmp, "Enable the 'Enable DXY Webhook Alerts' toggle in the Webhook / Alerts input group. "
     "Create ONE TradingView alert on the DXYTradeAlert indicator using the "
     "'alert() function calls' condition and enter your webhook URL.")

body(tmp, "DXY_ENTRY fires when a new DXY signal bar is confirmed. "
     "It tells your automation which direction is active on DXY and the key price levels.", bold=False)

sub_heading(tmp, "DXY_ENTRY JSON schema")
code_block(tmp, [
    '{',
    '  "event":      "DXY_ENTRY",',
    '  "signal":     "GAP_REJ_LONG",   // GAP_REJ_LONG/SHORT | REV_LONG/SHORT',
    '                                   // LON_ATTR_LONG | ATTR_CORE_LONG/SHORT',
    '  "dxy_dir":    "LONG",            // LONG or SHORT',
    '  "dxy_entry":  104.2500,          // DXY close at signal bar',
    '  "dxy_tp":     104.5000,          // DXY TP level',
    '  "dxy_sl":     104.0000,          // DXY SL level',
    '  "sl_pts":     2500,              // DXY SL distance in pts (÷10000 = price)',
    '  "gap_at_lon": 850,               // ATTR_CORE only — gap at London open (pts)',
    '  "wave_ext":   420,               // ATTR_CORE only — wave extension (pts)',
    '  "time":       "2024-01-15 09:00" // UTC bar close time',
    '}',
])

body(tmp, "DXY_EXIT fires when DXY hits its TP or SL level. "
     "This is the signal to close ALL open pair positions immediately.")

sub_heading(tmp, "DXY_EXIT JSON schema")
code_block(tmp, [
    '{',
    '  "event":    "DXY_EXIT",',
    '  "result":   "TP",                // TP or SL',
    '  "signal":   "GAP_REJ_LONG",      // signal that opened this trade',
    '  "dxy_exit": 104.4900,            // DXY close at exit bar',
    '  "time":     "2024-01-15 11:45"',
    '}',
])

body(tmp, "Note: DXYTradeAlert.pine also retains all original alertcondition() named alerts "
     "(push, email, pop-up). The new alert() JSON webhooks are additive — they fire in "
     "addition to, not instead of, the human-readable notifications.", italic=True)

# ── 12.3 DXYPairLevels.pine webhooks ─────────────────────────────────────────
sub_heading(tmp, "12.3  DXYPairLevels.pine — ENTRY, EXIT, and SL_MOVE events")

body(tmp, "Enable the 'Enable Webhook Alerts' toggle in the Webhook / Alerts group on each pair chart. "
     "Create ONE alert per chart using the 'alert() function calls' condition. "
     "All three event types are routed through this single alert.")

sub_heading(tmp, "ENTRY JSON schema")
code_block(tmp, [
    '{',
    '  "event":    "ENTRY",',
    '  "pair":     "EURUSD",',
    '  "signal":   "GAP_REJ_LONG",',
    '  "pair_dir": "SHORT",             // LONG or SHORT — the pair trade direction',
    '  "entry":    1.08500,             // pair close at signal bar',
    '  "tp":       1.08250,             // estimated pair TP (based on DXY SL distance × tick_factor)',
    '  "sl":       1.08750,             // estimated pair SL (1:1 mirror of TP distance)',
    '  "sl_pts":   2500,                // DXY SL distance in pts',
    '  "dxy_entry":104.2500,            // DXY close at signal bar',
    '  "gap_at_lon":850,                // ATTR_CORE only',
    '  "wave_ext": 420,                 // ATTR_CORE only',
    '  "time":     "2024-01-15 09:00"',
    '}',
])

sub_heading(tmp, "EXIT JSON schema")
code_block(tmp, [
    '{',
    '  "event":      "EXIT",',
    '  "pair":       "EURUSD",',
    '  "result":     "TP",              // TP or SL',
    '  "signal":     "GAP_REJ_LONG",',
    '  "exit_price": 1.08250,           // estimated pair exit price at DXY TP/SL bar',
    '  "dxy_exit":   104.4900,',
    '  "time":       "2024-01-15 11:45"',
    '}',
])

sub_heading(tmp, "SL_MOVE JSON schema  (new — breakeven advisory)")
code_block(tmp, [
    '{',
    '  "event":       "SL_MOVE",',
    '  "pair":        "EURUSD",',
    '  "signal":      "GAP_REJ_LONG",',
    '  "action":      "MOVE_TO_BE",     // always MOVE_TO_BE (breakeven)',
    '  "new_sl":      1.08500,          // new SL = pair entry price (breakeven)',
    '  "pair_entry":  1.08500,          // original pair entry price (same as new_sl)',
    '  "dxy_price":   104.3750,         // DXY price at advisory fire time',
    '  "progress_pct":50,               // % of DXY TP distance covered (default trigger: 50%)',
    '  "time":        "2024-01-15 10:15"',
    '}',
])

body(tmp, "SL_MOVE fires once per trade when DXY has moved the configured percentage "
     "of its TP distance from entry. The default trigger is 50% (halfway to TP). "
     "This is an advisory event only. NEVER act on SL_MOVE automatically — "
     "it notifies you to manually move your pair stop to breakeven.", bold=False)

body(tmp, "The SL_MOVE trigger percentage can be adjusted in the indicator settings "
     "(Webhook / Alerts group: 'Breakeven Move — Trigger'). Range: 25%–75%.")

# ── 12.4 TradingView alert setup ─────────────────────────────────────────────
sub_heading(tmp, "12.4  TradingView alert setup")

body(tmp, "For DXYTradeAlert.pine (DXY chart):", bold=True)
for step, txt in [
    ("1.", "Open the DXY chart with DXYTradeAlert applied."),
    ("2.", "In the indicator settings, open Webhook / Alerts and toggle ON 'Enable DXY Webhook Alerts'."),
    ("3.", "Click the bell icon → Create alert."),
    ("4.", "Condition: Select 'DXY Trade Alert' → 'alert() function calls'."),
    ("5.", "Notifications: Enable 'Webhook URL' and paste your endpoint."),
    ("6.", "Expiry: Set to 'Open-ended alert'. Save."),
]:
    bullet(tmp, txt, step + " ")

body(tmp, "For each DXYPairLevels.pine chart (repeat × number of pairs):", bold=True)
for step, txt in [
    ("1.", "Open the pair chart with DXYPairLevels applied."),
    ("2.", "In indicator settings, enable 'Enable Webhook Alerts' and optionally adjust 'Breakeven Move — Trigger'."),
    ("3.", "Click the bell icon → Create alert."),
    ("4.", "Condition: Select 'DXY Pair Levels' → 'alert() function calls'."),
    ("5.", "Notifications: Enable 'Webhook URL' — use the SAME endpoint as the DXY chart."),
    ("6.", "Expiry: Set to 'Open-ended alert'. Save."),
]:
    bullet(tmp, txt, step + " ")

body(tmp, "Result: all events (DXY_ENTRY, DXY_EXIT, ENTRY, EXIT, SL_MOVE) flow to a single "
     "endpoint. Route them inside your automation using the 'event' field.")

# ── 12.5 Automation routing logic ────────────────────────────────────────────
sub_heading(tmp, "12.5  Event routing logic")

body(tmp, "Always check the 'event' field first as the primary router:")

make_table(tmp,
    headers=["event", "Source", "Automation action", "Act automatically?"],
    rows=[
        ("DXY_ENTRY",  "DXYTradeAlert", "Log DXY signal active. Await per-pair ENTRY events.", ("Yes — log only", False, GREEN)),
        ("ENTRY",      "DXYPairLevels", "Open trade, log to journal. Set pair TP/SL from JSON.", ("Yes — execute trade", False, GREEN)),
        ("SL_MOVE",    "DXYPairLevels", "Notify trader to manually move stop to breakeven.", ("NO — advisory only", True, RGBColor(0xCC,0x33,0x00))),
        ("EXIT",       "DXYPairLevels", "Close pair position, log result to journal.", ("Yes — close trade", False, GREEN)),
        ("DXY_EXIT",   "DXYTradeAlert", "Emergency close-all if any pair positions still open.", ("Yes — safety net", False, AMBER)),
    ],
    col_widths=[3.0, 4.0, 6.0, 4.0]
)

body(tmp, "IMPORTANT: DXY_EXIT and per-pair EXIT events are redundant by design. "
     "In normal operation the pair EXIT fires first (same bar as DXY TP/SL). "
     "Use DXY_EXIT as a safety net in case a pair webhook was missed.", italic=True)

# ── 12.6 Platform guidance ────────────────────────────────────────────────────
sub_heading(tmp, "12.6  Receiving webhooks — platform guidance")

body(tmp, "TradingView sends webhook alerts as HTTP POST with the JSON string as the raw body "
     "(Content-Type: text/plain). Your endpoint must parse the raw body as JSON and "
     "respond with HTTP 200 within 10 seconds.")

make_table(tmp,
    headers=["Platform", "Use case", "Notes"],
    rows=[
        ("Make (Integromat)", "Journal to Google Sheets / Notion", "Use 'Custom Webhook' trigger. Map JSON fields to sheet columns. Best free-tier option."),
        ("n8n", "Full automation + broker API", "Self-hosted. Native JSON parsing. Can chain to broker REST API for auto-execution."),
        ("Zapier", "Simple journalling", "'Catch Hook' trigger. Limited JSON depth — may need intermediate step for nested fields."),
        ("AWS Lambda", "Custom broker integration", "Most flexible. Full control over execution logic. Requires coding."),
        ("Edgewonk / TraderSync", "Direct trade import", "Some SaaS journals accept webhook imports. Map fields at ingestion."),
    ],
    col_widths=[4.0, 5.0, 9.0]
)

# ── 12.7 Google Sheets journal example via Make ───────────────────────────────
sub_heading(tmp, "12.7  Example: Google Sheets journal via Make")

body(tmp, "This example logs ENTRY events to a Trades sheet and EXIT results to a Results sheet.")

for step, txt in [
    ("1.", "Create a Make scenario with a 'Custom Webhook' trigger. Copy the generated URL."),
    ("2.", "Paste the URL into all TradingView alert 'Webhook URL' fields (DXY chart + all pair charts)."),
    ("3.", "Add a Router module. Branch 1: filter event = 'ENTRY'. Branch 2: filter event = 'EXIT'. "
           "Branch 3: filter event = 'SL_MOVE'. Branch 4: filter event = 'DXY_ENTRY' or 'DXY_EXIT'."),
    ("4.", "Branch 1 — ENTRY → Google Sheets 'Add a Row' module on the Trades sheet.\n"
           "     Fields: pair→A, signal→B, pair_dir→C, entry→D, tp→E, sl→F, sl_pts→G, dxy_entry→H, time→I"),
    ("5.", "Branch 2 — EXIT → Google Sheets 'Search Rows' to find matching ENTRY by pair+signal+time approx, "
           "then 'Update a Row' with result, exit_price, dxy_exit."),
    ("6.", "Branch 3 — SL_MOVE → Send a push notification or Slack message: "
           "'Move [pair] stop to BE at [new_sl]. DXY [progress_pct]% toward TP.'"),
    ("7.", "Branch 4 — DXY_ENTRY/EXIT → Optional dashboard row or Slack channel message."),
]:
    bullet(tmp, txt, step + " ")

body(tmp, "IMPORTANT: Never place automatic orders from an SL_MOVE event. "
     "It is informational only. Manual confirmation before moving any stop is required.", bold=True)

# ── 12.8 Consistency checklist ────────────────────────────────────────────────
sub_heading(tmp, "12.8  Setup consistency checklist")

for item in [
    "DXYTradeAlert.pine: 'Enable DXY Webhook Alerts' toggle is ON on the DXY chart.",
    "DXYPairLevels.pine: 'Enable Webhook Alerts' toggle is ON on EVERY pair chart.",
    "One TradingView alert per chart — condition = 'alert() function calls'.",
    "All alerts point to the same webhook URL.",
    "Webhook endpoint responds HTTP 200 within 10 seconds.",
    "Automation routes on 'event' field FIRST before reading any other field.",
    "SL_MOVE is never acted on automatically — advisory only.",
    "DXY_EXIT used as safety net, not primary exit mechanism.",
    "Friday 19:30 UTC: DXYPairLevels fires a silent weekend reset (no webhook) — check for open positions manually.",
]:
    bullet(tmp, item)

# ─────────────────────────────────────────────────────────────────────────────
#  Splice the new Section 12 content into the main document
# ─────────────────────────────────────────────────────────────────────────────

# Collect elements from the temp document
tmp_elements = list(tmp.element.body)
# Remove the last element if it's an empty paragraph (Word always adds one)
while tmp_elements and tmp_elements[-1].tag.endswith("}p"):
    text = "".join(
        r.text for r in tmp_elements[-1].findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
    if text.strip() == "":
        tmp_elements.pop()
    else:
        break

# Insert before the Section 13 element
body_el = doc.element.body
all_elements_new = list(body_el)

if sec13_new_idx is not None:
    reference_elem = all_elements_new[sec13_new_idx]
    for elem in reversed(tmp_elements):
        elem_copy = copy.deepcopy(elem)
        reference_elem.addprevious(elem_copy)
else:
    # Append to end
    for elem in tmp_elements:
        body_el.append(copy.deepcopy(elem))

doc.save(str(DOCPATH))
print(f"Section 12 replaced successfully.")
print(f"Manual saved: {DOCPATH}")
